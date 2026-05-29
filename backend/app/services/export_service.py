import re
import html
import io
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def add_formatted_text(paragraph, html_text):
    """
    Parses simple inline HTML tags like <b>, <i>, <strong>, <em>
    and adds them as formatted runs to the paragraph.
    """
    if not html_text:
        return
    
    # Tokenize text by formatting tags
    tokens = re.split(r'(<b>|</b>|<strong>|</strong>|<i>|</i>|<em>|</em>|<br\s*/?>)', html_text)
    
    is_bold = False
    is_italic = False
    
    for token in tokens:
        if not token:
            continue
        if token in ['<b>', '<strong>']:
            is_bold = True
        elif token in ['</b>', '</strong>']:
            is_bold = False
        elif token in ['<i>', '<em>']:
            is_italic = True
        elif token in ['</i>', '</em>']:
            is_italic = False
        elif re.match(r'<br\s*/?>', token):
            paragraph.add_run('\n')
        else:
            # Strip remaining HTML tags and unescape symbols (e.g. &amp;, &lt;)
            clean_text = re.sub(r'<[^>]+>', '', token)
            clean_text = html.unescape(clean_text)
            
            run = paragraph.add_run(clean_text)
            run.bold = is_bold
            run.italic = is_italic

def write_editor_data_to_doc(doc, editor_data):
    """
    Parses EditorJS block structure and appends elements to the Word document.
    """
    if not editor_data or not isinstance(editor_data, dict):
        return
    
    blocks = editor_data.get("blocks", [])
    for block in blocks:
        b_type = block.get("type")
        b_data = block.get("data", {})
        
        if b_type == 'paragraph':
            text = b_data.get('text', '')
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            add_formatted_text(p, text)
            
        elif b_type == 'header':
            text = b_data.get('text', '')
            level = b_data.get('level', 2)
            # Standard Word heading styles can be used, but custom formatting is safer
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(html.unescape(text))
            run.bold = True
            if level == 1:
                run.font.size = Pt(14)
            elif level == 2:
                run.font.size = Pt(12)
            else:
                run.font.size = Pt(11)
                
        elif b_type == 'list':
            items = b_data.get('items', [])
            is_ordered = b_data.get('style') == 'ordered'
            
            for idx, item in enumerate(items):
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.left_indent = Inches(0.5)
                
                # Manual prefixes for lists to bypass Word's unpredictable list styles
                prefix = f"{idx + 1}.  " if is_ordered else "•  "
                p.add_run(prefix)
                add_formatted_text(p, item)
                
        elif b_type == 'math':
            latex = b_data.get('latex', '')
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.left_indent = Inches(0.25)
            run = p.add_run(f"$${latex}$$")
            run.italic = True
            
        elif b_type == 'image':
            # Images in EditorJS are URLs. Word export can skip images or add a placeholder text.
            caption = b_data.get('caption', '')
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(f"[IMAGE: {caption or 'Illustrative Figure'}]")
            run.bold = True
            run.font.size = Pt(10)
            
        elif b_type == 'table':
            content = b_data.get('content', [])
            if content:
                rows = len(content)
                cols = len(content[0]) if rows > 0 else 0
                if cols > 0:
                    table = doc.add_table(rows=rows, cols=cols)
                    table.style = 'Table Grid'
                    for r_idx, row_data in enumerate(content):
                        for c_idx, cell_value in enumerate(row_data):
                            cell = table.cell(r_idx, c_idx)
                            # Remove tags inside cells
                            clean_val = re.sub(r'<[^>]+>', '', cell_value)
                            cell.text = html.unescape(clean_val)
                    p_spacer = doc.add_paragraph()
                    p_spacer.paragraph_format.space_after = Pt(6)

def set_cell_border(cell, **kwargs):
    """
    Sets borders for table cells.
    kwargs can be top, bottom, left, right. Value is a dict with sz, val, color, space.
    e.g. set_cell_border(cell, bottom={"sz": 12, "val": "single", "color": "000000"})
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key, val in edge_data.items():
                element.set(qn('w:{}'.format(key)), str(val))

def generate_docx(paper):
    """
    Generates a structured Word document (.docx) matching standard university exam layouts.
    Returns a bytes buffer.
    """
    doc = Document()
    
    # 1. Page Margins (1 inch all around)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # 2. Base Typography
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    
    # 3. Header Title Block
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("QUESTION PAPER DESIGN SYSTEM (QPDS)\n")
    title_run.bold = True
    title_run.font.size = Pt(13)
    
    subtitle_run = title_p.add_run("Semester Examination Paper\n")
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(11)
    
    # 4. Meta Information Table
    subject = paper.subject
    ay_label = subject.academic_year.label if subject and subject.academic_year else "N/A"
    sem_number = f"Semester {subject.semester.number}" if subject and subject.semester else "N/A"
    sub_code = subject.code if subject else "N/A"
    sub_name = subject.name if subject else "N/A"
    paper_title = paper.title
    
    # Calculate Total Marks
    total_marks = 0
    for sec in paper.sections:
        for pq in sec.paper_questions:
            ed = pq.question.editor_data or {}
            marks_str = ed.get('marks') or ed.get('meta', {}).get('marks')
            try:
                total_marks += int(float(marks_str)) if marks_str else 0
            except:
                pass
                
    meta_table = doc.add_table(rows=3, cols=2)
    meta_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Fill details
    meta_table.cell(0, 0).text = f"Subject Name: {sub_name}"
    meta_table.cell(0, 1).text = f"Subject Code: {sub_code}"
    meta_table.cell(1, 0).text = f"Academic Year: {ay_label}"
    meta_table.cell(1, 1).text = f"Term/Level: {sem_number}"
    meta_table.cell(2, 0).text = f"Paper Title: {paper_title}"
    meta_table.cell(2, 1).text = f"Max Marks: {total_marks} Marks"
    
    # Border & Styling for Header Table
    for r in range(3):
        for c in range(2):
            cell = meta_table.cell(r, c)
            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            cell.paragraphs[0].runs[0].bold = True
            # Light border for layout structure
            set_cell_border(
                cell,
                top={"sz": 4, "val": "single", "color": "D3D3D3"},
                bottom={"sz": 4, "val": "single", "color": "D3D3D3"},
                left={"sz": 4, "val": "single", "color": "D3D3D3"},
                right={"sz": 4, "val": "single", "color": "D3D3D3"}
            )
            
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # Instructions Block
    inst_p = doc.add_paragraph()
    inst_p.paragraph_format.space_after = Pt(18)
    inst_run_title = inst_p.add_run("Instructions to Candidates:\n")
    inst_run_title.bold = True
    inst_run_title.font.size = Pt(10)
    
    inst_text = (
        "1. Ensure you have the complete question paper before writing.\n"
        "2. Answer all sections as per specific section instructions.\n"
        "3. Write clearly, displaying neat diagrams and equations where applicable."
    )
    inst_run_content = inst_p.add_run(inst_text)
    inst_run_content.font.size = Pt(9.5)
    inst_run_content.italic = True
    
    # Line separator
    sep_p = doc.add_paragraph()
    sep_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep_p.add_run("—" * 60)
    
    # 5. Populate Sections & Questions
    q_counter = 1
    for sec in paper.sections:
        sec_title_p = doc.add_paragraph()
        sec_title_p.paragraph_format.space_before = Pt(18)
        sec_title_p.paragraph_format.space_after = Pt(8)
        
        # Calculate section marks
        sec_marks = 0
        for pq in sec.paper_questions:
            ed = pq.question.editor_data or {}
            marks_str = ed.get('marks') or ed.get('meta', {}).get('marks')
            try:
                sec_marks += int(float(marks_str)) if marks_str else 0
            except:
                pass
                
        # Format Section Title, e.g., SECTION A: Questions [Total: 20 Marks]
        sec_title_text = f"{sec.title.upper()} ({sec_marks} Marks)"
        run_sec_title = sec_title_p.add_run(sec_title_text)
        run_sec_title.bold = True
        run_sec_title.font.size = Pt(12)
        run_sec_title.underline = True
        
        for pq in sec.paper_questions:
            question = pq.question
            ed = question.editor_data or {}
            marks_str = ed.get('marks') or ed.get('meta', {}).get('marks') or "0"
            co_code = question.course_outcome.co_code if question.course_outcome else ""
            
            # Question wrapper layout
            q_p = doc.add_paragraph()
            q_p.paragraph_format.space_before = Pt(8)
            q_p.paragraph_format.space_after = Pt(4)
            
            # Left-aligned question number
            prefix_run = q_p.add_run(f"Q{q_counter}.  ")
            prefix_run.bold = True
            
            # Extract main block text and format first block directly in Q paragraph
            blocks = ed.get('blocks', [])
            if blocks:
                first_block = blocks[0]
                if first_block.get('type') == 'paragraph':
                    add_formatted_text(q_p, first_block.get('data', {}).get('text', ''))
                else:
                    # If first block is list or table, parse it in subsequent blocks
                    # and leave Q1 header blank
                    pass
                    
            # Add metadata right-aligned style or inline
            meta_tag = f"   [{co_code}] ({marks_str} Marks)" if co_code else f"   ({marks_str} Marks)"
            meta_run = q_p.add_run(meta_tag)
            meta_run.bold = True
            
            # Print subsequent blocks (if any) under the question prefix
            if len(blocks) > 1:
                # Mock a document wrapper to write to the same doc
                sub_editor_data = {'blocks': blocks[1:]}
                write_editor_data_to_doc(doc, sub_editor_data)
                
            q_counter += 1
            
    # Save to dynamic memory stream
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


# ============================================================
# LaTeX Export
# ============================================================

def _escape_latex(text):
    """Escape characters that have special meaning in LaTeX."""
    if not text:
        return ''
    # Order matters: backslash first
    replacements = [
        ('\\', r'\textbackslash{}'),
        ('&', r'\&'),
        ('%', r'\%'),
        ('$', r'\$'),
        ('#', r'\#'),
        ('_', r'\_'),
        ('{', r'\{'),
        ('}', r'\}'),
        ('~', r'\textasciitilde{}'),
        ('^', r'\textasciicircum{}'),
    ]
    for old, new in replacements:
        text = text.replace(old, new)
    return text


def _html_to_latex(html_text):
    """
    Converts simple inline HTML formatting to LaTeX commands.
    Handles <b>, <strong>, <i>, <em>, <br>, and strips other tags.
    """
    if not html_text:
        return ''
    
    # First unescape HTML entities
    text = html.unescape(html_text)
    
    # Replace formatting tags with LaTeX equivalents
    # Process bold
    text = re.sub(r'<b>(.*?)</b>', r'\\textbf{\1}', text, flags=re.DOTALL)
    text = re.sub(r'<strong>(.*?)</strong>', r'\\textbf{\1}', text, flags=re.DOTALL)
    # Process italic
    text = re.sub(r'<i>(.*?)</i>', r'\\textit{\1}', text, flags=re.DOTALL)
    text = re.sub(r'<em>(.*?)</em>', r'\\textit{\1}', text, flags=re.DOTALL)
    # Line breaks
    text = re.sub(r'<br\s*/?>', r' \\\\ ', text)
    # Strip remaining HTML tags
    text = re.sub(r'<[^>]+>', '', text)
    
    # Escape LaTeX special chars in the remaining plain text portions
    # We need to be careful not to escape the LaTeX commands we just inserted
    # Strategy: split on our LaTeX commands, escape only the plain parts
    parts = re.split(r'(\\textbf\{.*?\}|\\textit\{.*?\}|\\\\)', text)
    result = []
    for part in parts:
        if part.startswith('\\textbf{') or part.startswith('\\textit{') or part == '\\\\':
            result.append(part)
        else:
            result.append(_escape_latex(part))
    
    return ''.join(result)


def _editor_blocks_to_latex(blocks):
    """Converts a list of EditorJS blocks into LaTeX body text."""
    lines = []
    
    for block in blocks:
        b_type = block.get('type')
        b_data = block.get('data', {})
        
        if b_type == 'paragraph':
            text = b_data.get('text', '')
            lines.append(_html_to_latex(text))
            lines.append('')  # blank line = paragraph break in LaTeX
            
        elif b_type == 'header':
            text = b_data.get('text', '')
            level = b_data.get('level', 2)
            clean = _escape_latex(html.unescape(text))
            if level == 1:
                lines.append(f'\\subsection*{{{clean}}}')
            elif level == 2:
                lines.append(f'\\subsubsection*{{{clean}}}')
            else:
                lines.append(f'\\paragraph{{{clean}}}')
            lines.append('')
            
        elif b_type == 'list':
            items = b_data.get('items', [])
            is_ordered = b_data.get('style') == 'ordered'
            env = 'enumerate' if is_ordered else 'itemize'
            lines.append(f'\\begin{{{env}}}')
            for item in items:
                lines.append(f'  \\item {_html_to_latex(item)}')
            lines.append(f'\\end{{{env}}}')
            lines.append('')
            
        elif b_type == 'math':
            latex_code = b_data.get('latex', '')
            if latex_code:
                lines.append('\\[')
                lines.append(f'  {latex_code}')
                lines.append('\\]')
                lines.append('')
                
        elif b_type == 'image':
            caption = b_data.get('caption', '')
            clean_caption = _escape_latex(html.unescape(caption)) if caption else 'Illustrative Figure'
            lines.append(f'\\begin{{center}}')
            lines.append(f'  [IMAGE: {clean_caption}]')
            lines.append(f'\\end{{center}}')
            lines.append('')
            
        elif b_type == 'table':
            content = b_data.get('content', [])
            if content and len(content) > 0:
                cols = len(content[0])
                col_spec = '|'.join(['l'] * cols)
                lines.append(f'\\begin{{tabular}}{{|{col_spec}|}}')
                lines.append('\\hline')
                for row_data in content:
                    cells = [_escape_latex(html.unescape(re.sub(r'<[^>]+>', '', c))) for c in row_data]
                    lines.append(' & '.join(cells) + ' \\\\')
                    lines.append('\\hline')
                lines.append('\\end{tabular}')
                lines.append('')
    
    return '\n'.join(lines)


def generate_latex(paper):
    """
    Generates a compileable LaTeX (.tex) document from a Paper object.
    Returns a BytesIO stream containing the .tex file content.
    """
    subject = paper.subject
    ay_label = subject.academic_year.label if subject and subject.academic_year else 'N/A'
    sem_number = f'Semester {subject.semester.number}' if subject and subject.semester else 'N/A'
    sub_code = _escape_latex(subject.code) if subject else 'N/A'
    sub_name = _escape_latex(subject.name) if subject else 'N/A'
    paper_title = _escape_latex(paper.title)
    
    # Calculate Total Marks
    total_marks = 0
    for sec in paper.sections:
        for pq in sec.paper_questions:
            ed = pq.question.editor_data or {}
            marks_str = ed.get('marks') or ed.get('meta', {}).get('marks')
            try:
                total_marks += int(float(marks_str)) if marks_str else 0
            except:
                pass
    
    # Build the document
    lines = []
    
    # Preamble
    lines.append(r'\documentclass[12pt,a4paper]{article}')
    lines.append(r'\usepackage[margin=1in]{geometry}')
    lines.append(r'\usepackage{amsmath,amssymb}')
    lines.append(r'\usepackage{enumitem}')
    lines.append(r'\usepackage{tabularx}')
    lines.append(r'\usepackage{graphicx}')
    lines.append(r'\usepackage{fancyhdr}')
    lines.append(r'\usepackage[utf8]{inputenc}')
    lines.append(r'\usepackage[T1]{fontenc}')
    lines.append(r'\usepackage{times}')
    lines.append('')
    lines.append(r'\pagestyle{fancy}')
    lines.append(r'\fancyhf{}')
    lines.append(r'\rfoot{\thepage}')
    lines.append(r'\renewcommand{\headrulewidth}{0pt}')
    lines.append('')
    lines.append(r'\begin{document}')
    lines.append('')
    
    # Title Block
    lines.append(r'\begin{center}')
    lines.append(r'  {\Large \textbf{QUESTION PAPER DESIGN SYSTEM (QPDS)}} \\[6pt]')
    lines.append(r'  {\large \textbf{Semester Examination Paper}} \\[12pt]')
    lines.append(r'\end{center}')
    lines.append('')
    
    # Meta Information Table
    lines.append(r'\noindent')
    lines.append(r'\begin{tabularx}{\textwidth}{|X|X|}')
    lines.append(r'\hline')
    lines.append(f'\\textbf{{Subject Name:}} {sub_name} & \\textbf{{Subject Code:}} {sub_code} \\\\')
    lines.append(r'\hline')
    lines.append(f'\\textbf{{Academic Year:}} {_escape_latex(ay_label)} & \\textbf{{Term/Level:}} {_escape_latex(sem_number)} \\\\')
    lines.append(r'\hline')
    lines.append(f'\\textbf{{Paper Title:}} {paper_title} & \\textbf{{Max Marks:}} {total_marks} Marks \\\\')
    lines.append(r'\hline')
    lines.append(r'\end{tabularx}')
    lines.append('')
    lines.append(r'\vspace{12pt}')
    lines.append('')
    
    # Instructions
    lines.append(r'\noindent\textbf{Instructions to Candidates:}')
    lines.append(r'\begin{enumerate}[nosep]')
    lines.append(r'  \item Ensure you have the complete question paper before writing.')
    lines.append(r'  \item Answer all sections as per specific section instructions.')
    lines.append(r'  \item Write clearly, displaying neat diagrams and equations where applicable.')
    lines.append(r'\end{enumerate}')
    lines.append('')
    lines.append(r'\noindent\rule{\textwidth}{0.4pt}')
    lines.append(r'\vspace{12pt}')
    lines.append('')
    
    # Sections & Questions
    q_counter = 1
    for sec in paper.sections:
        # Calculate section marks
        sec_marks = 0
        for pq in sec.paper_questions:
            ed = pq.question.editor_data or {}
            marks_str = ed.get('marks') or ed.get('meta', {}).get('marks')
            try:
                sec_marks += int(float(marks_str)) if marks_str else 0
            except:
                pass
        
        sec_title = _escape_latex(sec.title.upper())
        lines.append(f'\\section*{{\\underline{{{sec_title} ({sec_marks} Marks)}}}}')
        lines.append(r'\vspace{6pt}')
        lines.append('')
        
        for pq in sec.paper_questions:
            question = pq.question
            ed = question.editor_data or {}
            marks_str = ed.get('marks') or ed.get('meta', {}).get('marks') or '0'
            co_code = question.course_outcome.co_code if question.course_outcome else ''
            
            # Question number and marks tag
            marks_tag = f'[{_escape_latex(co_code)}] ({marks_str} Marks)' if co_code else f'({marks_str} Marks)'
            
            blocks = ed.get('blocks', [])
            
            # First block inline with question number
            first_text = ''
            if blocks:
                first_block = blocks[0]
                if first_block.get('type') == 'paragraph':
                    first_text = _html_to_latex(first_block.get('data', {}).get('text', ''))
            
            lines.append(f'\\noindent\\textbf{{Q{q_counter}.}} {first_text} \\hfill \\textbf{{{marks_tag}}}')
            lines.append('')
            
            # Subsequent blocks
            if len(blocks) > 1:
                latex_body = _editor_blocks_to_latex(blocks[1:])
                if latex_body.strip():
                    lines.append(latex_body)
            
            lines.append(r'\vspace{8pt}')
            lines.append('')
            q_counter += 1
    
    lines.append(r'\end{document}')
    
    # Write to buffer
    content = '\n'.join(lines)
    file_stream = io.BytesIO(content.encode('utf-8'))
    file_stream.seek(0)
    return file_stream

